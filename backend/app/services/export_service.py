# -*- coding: utf-8 -*-
"""Export generation service.

S8 扩展：在原有 Excel/bundle（legacy ReviewMatch 链路）基础上新增
- 报告导出（pdf=reportlab / docx=python-docx / html=模板字符串），
  数据源为 S7 章节化报告（ReportChapter + 可选 ReportAnnotation）.
- 数据导出（excel/csv），数据源为 S5 flow_records（raw/standard）或 S6 findings.
- 导出历史列表 + 预览取样（不生成产物）.

不删减精神：导出只读原数据 + 复制产物，不删原记录；导出历史产物文件保留可
重新下载（不删 ExportFile 行 / 不删产物文件）.
"""

import csv
import json
import zipfile
from pathlib import Path
from typing import Any

import openpyxl
from openpyxl.styles import Font, PatternFill
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import settings
from app.models import (
    ExportFile,
    Finding,
    FlowRecordRow,
    Report,
    ReportAnnotation,
    ReportChapter,
    Review,
    ReviewMatch,
    Task,
)
from app.services.review_service import FlowRecord, ReviewService


class ExportService:
    """Generate Excel and skills-bundle artifacts for a task."""

    def __init__(self, output_dir: str | None = None):
        self.output_dir = Path(output_dir or settings.OUTPUT_DIR) / "exports"
        self.review_service = ReviewService()

    async def export_excel(self, db: AsyncSession, task_id: int, review_id: int | None = None) -> ExportFile:
        """Export standardized records and matches to an Excel workbook."""
        task = await self._load_task(db, task_id)
        review = await self._load_review(db, task_id, review_id)
        records = await self.review_service.load_task_records(db, task_id)
        matches = await self._load_matches(db, review.id if review else None)

        export_dir = self.output_dir / str(task_id)
        export_dir.mkdir(parents=True, exist_ok=True)
        path = export_dir / f"task_{task_id}_review.xlsx"

        wb = openpyxl.Workbook()
        try:
            self._write_records_sheet(wb.active, records, matches)
            self._write_matches_sheet(wb, matches)
            wb.save(path)
        finally:
            wb.close()

        export = ExportFile(
            task_id=task.id,
            review_id=review.id if review else None,
            format="excel",
            file_path=str(path),
        )
        db.add(export)
        await db.flush()
        await db.refresh(export)
        return export

    async def export_bundle(self, db: AsyncSession, task_id: int, review_id: int | None = None) -> ExportFile:
        """Export a minimal task skills bundle ZIP."""
        task = await self._load_task(db, task_id)
        review = await self._load_review(db, task_id, review_id)
        matches = await self._load_matches(db, review.id if review else None)

        export_dir = self.output_dir / str(task_id)
        export_dir.mkdir(parents=True, exist_ok=True)
        path = export_dir / f"task_{task_id}_skills_bundle.zip"

        manifest = {
            "bundle_type": "employee_customer_audit_skill",
            "bundle_version": "1.0.0",
            "task_id": task.id,
            "task_title": task.title,
            "review_id": review.id if review else None,
            "match_count": len(matches),
        }
        review_payload = {
            "task": {"id": task.id, "title": task.title, "status": task.status},
            "review": {"id": review.id, "status": review.status} if review else None,
            "matches": [self._match_to_dict(match) for match in matches],
        }

        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("skill_manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
            zf.writestr("current_task/review_result.json", json.dumps(review_payload, ensure_ascii=False, indent=2))
            zf.writestr("SKILL.md", self._skill_markdown(task, len(matches)))

        export = ExportFile(
            task_id=task.id,
            review_id=review.id if review else None,
            format="bundle",
            file_path=str(path),
        )
        db.add(export)
        await db.flush()
        await db.refresh(export)
        return export

    # -------------------------------------------------------------------
    # S8 报告导出：pdf / docx / html
    # -------------------------------------------------------------------

    async def export_report(
        self,
        db: AsyncSession,
        task_id: int,
        fmt: str,
        include_annotations: bool = False,
    ) -> ExportFile:
        """基于 S7 章节化报告生成 pdf/docx/html 产物.

        数据源：task 最新 Report + chapters（按 order_index 排序）+
        可选 annotations. 报告不存在 → ValueError（router 转 404）.
        """
        task = await self._load_task(db, task_id)
        report = await self._load_task_report(db, task.id)
        if report.status == "generating":
            raise ValueError("Report is still generating")
        if report.status == "failed":
            raise ValueError("Report generation failed")
        chapters = sorted(report.chapters, key=lambda c: c.order_index)
        annotations = list(report.annotations) if include_annotations else []

        export_dir = self.output_dir / str(task.id)
        export_dir.mkdir(parents=True, exist_ok=True)
        ext = {"pdf": "pdf", "docx": "docx", "html": "html"}.get(fmt, fmt)
        path = export_dir / f"task_{task.id}_report_{fmt}.{ext}"

        if fmt == "pdf":
            self._write_report_pdf(path, task, chapters, annotations)
        elif fmt == "docx":
            self._write_report_docx(path, task, chapters, annotations)
        elif fmt == "html":
            self._write_report_html(path, task, chapters, annotations)
        else:
            raise ValueError("Unsupported report format: %s" % fmt)

        export = ExportFile(
            task_id=task.id,
            review_id=report.review_id,
            format=fmt,
            scope="report",
            file_path=str(path),
        )
        db.add(export)
        await db.flush()
        await db.refresh(export)
        return export

    # -------------------------------------------------------------------
    # S8 数据导出：excel / csv × raw / standard / findings
    # -------------------------------------------------------------------

    async def export_data(
        self,
        db: AsyncSession,
        task_id: int,
        scope: str,
        fmt: str,
    ) -> ExportFile:
        """导出 flow_records（raw/standard）或 findings 为 excel/csv."""
        task = await self._load_task(db, task_id)

        export_dir = self.output_dir / str(task.id)
        export_dir.mkdir(parents=True, exist_ok=True)
        ext = "xlsx" if fmt == "excel" else "csv"
        path = export_dir / f"task_{task.id}_{scope}.{ext}"

        if scope in ("raw", "standard"):
            rows = await self._load_flow_records(db, task.id, scope)
            if fmt == "excel":
                self._write_flow_records_excel(path, scope, rows)
            else:
                self._write_flow_records_csv(path, scope, rows)
        elif scope == "findings":
            rows = await self._load_findings(db, task.id)
            if fmt == "excel":
                self._write_findings_excel(path, rows)
            else:
                self._write_findings_csv(path, rows)
        else:
            raise ValueError("Unsupported data scope: %s" % scope)

        export = ExportFile(
            task_id=task.id,
            format=fmt,
            scope=scope,
            file_path=str(path),
        )
        db.add(export)
        await db.flush()
        await db.refresh(export)
        return export

    async def list_task_exports(self, db: AsyncSession, task_id: int) -> list[ExportFile]:
        """导出历史列表（按 created_at 降序，含 scope/format/file_path/created_at）."""
        result = await db.execute(
            select(ExportFile)
            .where(ExportFile.task_id == task_id)
            .order_by(ExportFile.created_at.desc())
        )
        return list(result.scalars().all())

    async def preview_export(
        self,
        db: AsyncSession,
        task_id: int,
        scope: str,
    ) -> dict[str, Any]:
        """取样预览（不生成产物）.

        - report: 前 2 章 content 文本 + 批注数.
        - raw / standard: 前 20 行 JSON.
        - findings: 前 20 行 JSON.
        """
        task = await self._load_task(db, task_id)
        if scope == "report":
            report = await self._load_task_report(db, task.id)
            chapters = sorted(report.chapters, key=lambda c: c.order_index)
            sample = [
                {"title": c.title, "content": c.content}
                for c in chapters[:2]
            ]
            return {
                "scope": "report",
                "sample": sample,
                "annotation_count": len(report.annotations),
            }
        if scope in ("raw", "standard"):
            rows = await self._load_flow_records(db, task.id, scope)
            return {
                "scope": scope,
                "sample": [self._flow_record_to_dict(r) for r in rows[:20]],
            }
        if scope == "findings":
            rows = await self._load_findings(db, task.id)
            return {
                "scope": "findings",
                "sample": [self._finding_to_dict(r) for r in rows[:20]],
            }
        raise ValueError("Unsupported preview scope: %s" % scope)

    # -------------------------------------------------------------------
    # S8 helpers: data loading
    # -------------------------------------------------------------------

    async def _load_task_report(self, db: AsyncSession, task_id: int) -> Report:
        """取 task 最新章节化报告 + chapters + annotations（不存在 raise ValueError → 404）."""
        result = await db.execute(
            select(Report)
            .options(
                selectinload(Report.chapters),
                selectinload(Report.annotations),
            )
            .where(Report.task_id == task_id)
            .order_by(Report.created_at.desc())
            .limit(1)
        )
        report = result.scalar_one_or_none()
        if not report:
            raise ValueError("Report not generated yet")
        return report

    async def _load_flow_records(
        self, db: AsyncSession, task_id: int, scope: str
    ) -> list[FlowRecordRow]:
        """raw=全部 flow_records；standard=record_type='standard'."""
        query = select(FlowRecordRow).where(FlowRecordRow.task_id == task_id)
        if scope == "standard":
            query = query.where(FlowRecordRow.record_type == "standard")
        query = query.order_by(FlowRecordRow.id.asc())
        result = await db.execute(query)
        return list(result.scalars().all())

    async def _load_findings(self, db: AsyncSession, task_id: int) -> list[Finding]:
        """S6 findings 全部（按 id 升序）."""
        result = await db.execute(
            select(Finding)
            .where(Finding.task_id == task_id)
            .order_by(Finding.id.asc())
        )
        return list(result.scalars().all())

    # -------------------------------------------------------------------
    # S8 helpers: report writers (pdf / docx / html) — 黑白单色年报排版
    # -------------------------------------------------------------------

    def _write_report_pdf(
        self,
        path: Path,
        task: Task,
        chapters: list[ReportChapter],
        annotations: list[ReportAnnotation],
    ) -> None:
        """reportlab platypus 黑白年报排版（封面 + 目录 + 各章 + 批注附录）.

        结构：封面整页 → 目录页（带真实页码 + leader 点）→ 正文各章 → 批注附录。
        目录用 reportlab TableOfContents + 自定义 doc template（afterFlowable
        在每章标题渲染后发 ``TOCEntry`` 通知，multiBuild 两遍渲染拿真实页码）。
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.lib.enums import TA_CENTER
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            PageBreak,
            HRFlowable,
        )
        from reportlab.platypus.tableofcontents import TableOfContents

        # 章标题 style 名（afterFlowable 据此识别并发 TOCEntry）。
        CHAPTER_HEADING_STYLE = "ChapterHeading"

        class TocDocTemplate(SimpleDocTemplate):
            """自定义 doc template：章标题渲染后通知 TOC（拿真实页码）."""

            def afterFlowable(self, flowable):
                if isinstance(flowable, Paragraph):
                    style_name = flowable.style.name
                    if style_name == CHAPTER_HEADING_STYLE:
                        text = flowable.getPlainText()
                        page = self.page
                        self.notify("TOCEntry", (0, text, page))

        doc = TocDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
            title=f"审查报告 - {task.title}",
        )
        base = getSampleStyleSheet()
        # 封面样式（居中大标题 + 副标题 + 元信息）。
        cover_title_style = ParagraphStyle(
            "CoverTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=28,
            leading=36,
            alignment=TA_CENTER,
            textColor="#000000",
        )
        cover_sub_style = ParagraphStyle(
            "CoverSub",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=14,
            leading=20,
            alignment=TA_CENTER,
            textColor="#1f1f1f",
            spaceBefore=8,
            spaceAfter=8,
        )
        cover_meta_style = ParagraphStyle(
            "CoverMeta",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=11,
            leading=18,
            alignment=TA_CENTER,
            textColor="#595959",
        )
        # 目录标题样式（不用 ChapterHeading，避免目录页自身进 TOC）。
        toc_title_style = ParagraphStyle(
            "TOCTitle",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=26,
            alignment=TA_CENTER,
            textColor="#000000",
            spaceAfter=16,
        )
        # 章标题样式（afterFlowable 据此 style 名发 TOCEntry）。
        chapter_heading_style = ParagraphStyle(
            CHAPTER_HEADING_STYLE,
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=24,
            spaceBefore=14,
            spaceAfter=8,
            textColor="#000000",
        )
        body_style = ParagraphStyle(
            "ReportBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=16,
            spaceAfter=6,
            textColor="#1f1f1f",
        )
        ann_style = ParagraphStyle(
            "ReportAnn",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=14,
            leftIndent=12,
            textColor="#595959",
        )

        toc = TableOfContents()
        toc.levelStyles = [
            ParagraphStyle(
                name="TOCLevel1",
                fontName="Helvetica",
                fontSize=11,
                leading=20,
                leftIndent=20,
                firstLineIndent=-20,
                spaceBefore=0,
                textColor="#1f1f1f",
            )
        ]

        flowables: list[Any] = []

        # ----- 封面（整页）-----
        flowables.append(Spacer(1, 50 * mm))
        flowables.append(HRFlowable(width="100%", thickness=1, color="#000000"))
        flowables.append(Spacer(1, 16 * mm))
        flowables.append(Paragraph("银行/支付流水审查报告", cover_title_style))
        flowables.append(Spacer(1, 10 * mm))
        flowables.append(Paragraph(self._escape_html(task.title or ""), cover_sub_style))
        flowables.append(Spacer(1, 24 * mm))

        # 元信息块（缺项跳过）。
        meta_lines = self._cover_meta_lines(task)
        for line in meta_lines:
            flowables.append(Paragraph(self._escape_html(line), cover_meta_style))
        flowables.append(Spacer(1, 16 * mm))
        flowables.append(HRFlowable(width="100%", thickness=1, color="#000000"))
        flowables.append(PageBreak())

        # ----- 目录页 -----
        flowables.append(Paragraph("目录", toc_title_style))
        flowables.append(toc)
        flowables.append(PageBreak())

        # ----- 正文各章 -----
        # 章标题用 chapter_heading_style（afterFlowable 发 TOCEntry 抓真实页码）；
        # 章节正文走 markdown 渲染器（report_markdown.render_pdf）正确排版
        # ##/列表/表格/加粗，不再字面显示符号。
        from app.services.report_markdown import parse_markdown_blocks, render_pdf

        md_styles = {
            "h2": ParagraphStyle(
                "MdH2",
                parent=base["Heading2"],
                fontName="Helvetica-Bold",
                fontSize=14,
                leading=18,
                spaceBefore=10,
                spaceAfter=4,
                textColor="#000000",
            ),
            "h3": ParagraphStyle(
                "MdH3",
                parent=base["Heading3"],
                fontName="Helvetica-Bold",
                fontSize=12,
                leading=16,
                spaceBefore=8,
                spaceAfter=3,
                textColor="#000000",
            ),
            "body": body_style,
            "quote": ParagraphStyle(
                "MdQuote",
                parent=base["BodyText"],
                fontName="Helvetica",
                fontSize=10.5,
                leading=15,
                leftIndent=12,
                spaceAfter=6,
                textColor="#595959",
            ),
        }
        for ch in chapters:
            flowables.append(Paragraph(ch.title, chapter_heading_style))
            render_pdf(flowables, parse_markdown_blocks(ch.content), md_styles)
            flowables.append(Spacer(1, 3 * mm))

        if annotations:
            flowables.append(PageBreak())
            flowables.append(Paragraph("批注附录", chapter_heading_style))
            for ann in annotations:
                label = "已解决" if ann.resolved else "待解决"
                head = f"[{label}] {ann.author}："
                flowables.append(Paragraph(self._escape_html(head + ann.content), ann_style))

        # multiBuild 两遍渲染：第一遍拿页码填 TOC，第二遍生成最终页。
        doc.multiBuild(flowables)

    def _write_report_docx(
        self,
        path: Path,
        task: Task,
        chapters: list[ReportChapter],
        annotations: list[ReportAnnotation],
    ) -> None:
        """python-docx 黑白年报排版（封面整页 + 目录域 + 各章 + 批注附录）.

        结构：封面整页（大标题/副标题/元信息块/横线）→ 目录页（原生 TOC 域
        ``TOC \\o "1-2"``，用户在 Word 里「更新域」才显页码）→ 正文各章 →
        批注附录。封面/目录后各插 page break。
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt, RGBColor

        doc = Document()

        # ----- 封面整页 -----
        # 上横线。
        self._docx_horizontal_rule(doc)
        # 顶部留白。
        for _ in range(4):
            doc.add_paragraph("")
        # 大标题（居中、加粗、大字号）。
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run("银行/支付流水审查报告")
        run.bold = True
        run.font.size = Pt(32)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        doc.add_paragraph("")
        # 副标题（任务名，居中）。
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(task.title or "")
        sub_run.font.size = Pt(16)
        sub_run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        # 留白到中下部。
        for _ in range(6):
            doc.add_paragraph("")
        # 元信息块（居中、缺项跳过）。
        meta_lines = self._cover_meta_lines(task)
        for line in meta_lines:
            meta_p = doc.add_paragraph()
            meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta_run = meta_p.add_run(line)
            meta_run.font.size = Pt(11)
            meta_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        # 留白。
        for _ in range(3):
            doc.add_paragraph("")
        # 下横线。
        self._docx_horizontal_rule(doc)
        # 封面后分页。
        doc.add_page_break()

        # ----- 目录页（原生 TOC 域）-----
        toc_title_p = doc.add_paragraph()
        toc_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title_run = toc_title_p.add_run("目录")
        toc_title_run.bold = True
        toc_title_run.font.size = Pt(20)
        toc_title_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        doc.add_paragraph("")
        # 插入 TOC 域：TOC \o "1-2"（捕获 Heading 1-2）。
        self._docx_insert_toc_field(doc)
        # 提示文案。
        hint_p = doc.add_paragraph()
        hint_run = hint_p.add_run("打开后若未显示页码，请右键→更新域。")
        hint_run.font.size = Pt(9)
        hint_run.italic = True
        hint_run.font.color.rgb = RGBColor(0x8C, 0x8C, 0x8C)
        # 目录后分页。
        doc.add_page_break()

        # ----- 正文各章 -----
        # 章标题用 Heading 1（TOC 域 \o "1-2" 捕获）；章节正文走 markdown 渲染器
        # （report_markdown.render_docx）正确排版 ##/列表/表格/加粗，不再字面显示符号。
        from app.services.report_markdown import parse_markdown_blocks, render_docx

        for ch in chapters:
            doc.add_heading(ch.title, level=1)
            render_docx(doc, parse_markdown_blocks(ch.content))

        if annotations:
            doc.add_heading("批注附录", level=1)
            for ann in annotations:
                label = "已解决" if ann.resolved else "待解决"
                p = doc.add_paragraph(f"[{label}] {ann.author}：{ann.content}")
                for run in p.runs:
                    run.font.size = Pt(9.5)

        doc.save(str(path))

    def _write_report_html(
        self,
        path: Path,
        task: Task,
        chapters: list[ReportChapter],
        annotations: list[ReportAnnotation],
    ) -> None:
        """模板字符串生成自包含 HTML（单色内联 CSS，黑白年报）.

        含封面（大标题/副标题/元信息块/横线）+ 各章 + 批注附录。
        """
        parts: list[str] = [
            "<!DOCTYPE html>",
            '<html lang="zh-CN">',
            "<head>",
            '<meta charset="utf-8">',
            f"<title>审查报告 - {self._escape_html(task.title)}</title>",
            "<style>",
            "body{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;"
            "color:#1f1f1f;background:#fff;max-width:800px;margin:24px auto;padding:0 16px;}",
            ".cover{text-align:center;padding:60px 0 40px;border-bottom:1px solid #000;margin-bottom:32px;}",
            ".cover h1{font-size:32px;font-weight:700;border:none;padding:0;margin-bottom:12px;}",
            ".cover .sub{font-size:16px;color:#1f1f1f;margin-bottom:32px;}",
            ".cover .meta{color:#595959;font-size:13px;margin:6px 0;}",
            "h1{font-size:28px;font-weight:700;border-bottom:1px solid #000;padding-bottom:8px;}",
            "h2{font-size:20px;font-weight:700;margin-top:24px;}",
            "p{font-size:14px;line-height:1.7;white-space:pre-wrap;}",
            ".meta{color:#595959;font-size:13px;margin:4px 0;}",
            ".ann{border-left:2px solid #bfbfbf;background:#f0f0f0;padding:8px 12px;margin:8px 0;"
            "font-size:13px;color:#595959;}",
            ".ann-label{font-weight:700;color:#000;}",
            "</style>",
            "</head>",
            "<body>",
            '<div class="cover">',
            "<h1>银行/支付流水审查报告</h1>",
            f'<div class="sub">{self._escape_html(task.title or "")}</div>',
        ]
        for line in self._cover_meta_lines(task):
            parts.append(f'<p class="meta">{self._escape_html(line)}</p>')
        parts.append("</div>")
        for ch in chapters:
            parts.append(f"<h2>{self._escape_html(ch.title)}</h2>")
            for html_block in self._markdown_blocks_to_html(ch.content):
                parts.append(html_block)
        if annotations:
            parts.append("<h2>批注附录</h2>")
            for ann in annotations:
                label = "已解决" if ann.resolved else "待解决"
                parts.append(
                    f'<div class="ann"><span class="ann-label">[{label}] '
                    f"{self._escape_html(ann.author)}：</span>"
                    f"{self._escape_html(ann.content)}</div>"
                )
        parts.append("</body></html>")
        path.write_text("\n".join(parts), encoding="utf-8")

    # -------------------------------------------------------------------
    # S8 helpers: data writers (excel / csv) — 黑白表头
    # -------------------------------------------------------------------

    def _write_flow_records_excel(
        self, path: Path, scope: str, rows: list[FlowRecordRow]
    ) -> None:
        """flow_records excel（黑白表头，raw 含 raw_payload 序列化列）."""
        wb = openpyxl.Workbook()
        try:
            ws = wb.active
            ws.title = f"{scope}_records"
            headers = self._flow_record_headers(scope)
            self._write_header(ws, headers)
            for row_index, rec in enumerate(rows, 2):
                values = self._flow_record_row(scope, rec)
                for col_index, value in enumerate(values, 1):
                    ws.cell(row=row_index, column=col_index, value=value)
            wb.save(path)
        finally:
            wb.close()

    def _write_flow_records_csv(
        self, path: Path, scope: str, rows: list[FlowRecordRow]
    ) -> None:
        """flow_records csv（UTF-8 BOM，raw 含 raw_payload JSON）."""
        headers = self._flow_record_headers(scope)
        with open(path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            for rec in rows:
                writer.writerow(self._flow_record_row(scope, rec))

    def _write_findings_excel(self, path: Path, rows: list[Finding]) -> None:
        """findings excel（黑白表头）."""
        wb = openpyxl.Workbook()
        try:
            ws = wb.active
            ws.title = "findings"
            headers = self._finding_headers()
            self._write_header(ws, headers)
            for row_index, finding in enumerate(rows, 2):
                values = self._finding_row(finding)
                for col_index, value in enumerate(values, 1):
                    ws.cell(row=row_index, column=col_index, value=value)
            wb.save(path)
        finally:
            wb.close()

    def _write_findings_csv(self, path: Path, rows: list[Finding]) -> None:
        """findings csv（UTF-8 BOM）."""
        headers = self._finding_headers()
        with open(path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            for finding in rows:
                writer.writerow(self._finding_row(finding))

    @staticmethod
    def _flow_record_headers(scope: str) -> list[str]:
        base = [
            "记录ID",
            "文档ID",
            "渠道",
            "记录类型",
            "行号",
            "交易时间",
            "交易对手名",
            "交易对手账号",
            "金额",
            "原始金额",
            "摘要",
            "收支类型",
            "状态",
            "排除原因",
        ]
        if scope == "raw":
            base.append("原始载荷")
        return base

    @staticmethod
    def _flow_record_row(scope: str, rec: FlowRecordRow) -> list[Any]:
        row: list[Any] = [
            rec.id,
            rec.document_id,
            rec.channel or "",
            rec.record_type,
            rec.row_index,
            rec.transaction_time or "",
            rec.counterparty_name or "",
            rec.counterparty_account or "",
            rec.amount or "",
            rec.raw_amount or "",
            rec.summary or "",
            rec.transaction_type or "",
            rec.status,
            rec.exclude_reason or "",
        ]
        if scope == "raw":
            payload = rec.raw_payload if rec.raw_payload is not None else ""
            row.append(json.dumps(payload, ensure_ascii=False) if payload else "")
        return row

    @staticmethod
    def _flow_record_to_dict(rec: FlowRecordRow) -> dict[str, Any]:
        return {
            "id": rec.id,
            "document_id": rec.document_id,
            "channel": rec.channel,
            "record_type": rec.record_type,
            "row_index": rec.row_index,
            "transaction_time": rec.transaction_time,
            "counterparty_name": rec.counterparty_name,
            "counterparty_account": rec.counterparty_account,
            "amount": rec.amount,
            "raw_amount": rec.raw_amount,
            "summary": rec.summary,
            "transaction_type": rec.transaction_type,
            "status": rec.status,
            "exclude_reason": rec.exclude_reason,
            "raw_payload": rec.raw_payload,
        }

    @staticmethod
    def _finding_headers() -> list[str]:
        return [
            "发现ID",
            "类型",
            "风险等级",
            "描述",
            "交易对手",
            "金额",
            "置信度",
            "状态",
            "备注",
            "创建时间",
        ]

    @staticmethod
    def _finding_row(f: Finding) -> list[Any]:
        return [
            f.id,
            f.type,
            f.severity,
            f.description,
            f.counterparty or "",
            f.amount or "",
            f.confidence,
            f.status,
            f.comment or "",
            f.created_at.isoformat() if f.created_at else "",
        ]

    @staticmethod
    def _finding_to_dict(f: Finding) -> dict[str, Any]:
        return {
            "id": f.id,
            "type": f.type,
            "severity": f.severity,
            "description": f.description,
            "counterparty": f.counterparty,
            "amount": f.amount,
            "confidence": f.confidence,
            "status": f.status,
            "comment": f.comment,
            "created_at": f.created_at.isoformat() if f.created_at else "",
            "updated_at": f.updated_at.isoformat() if f.updated_at else "",
        }

    def _markdown_blocks_to_html(self, md: str) -> list[str]:
        """把约束子集 markdown 渲染成 HTML 片段（用于 _write_report_html）.

        复用 report_markdown.parse_markdown_blocks 解析，转成 HTML 字符串列表。
        单色：表头灰底加粗（inline style），引用灰字。
        """
        from app.services.report_markdown import parse_markdown_blocks

        parts: list[str] = []
        for block in parse_markdown_blocks(md):
            btype = block["type"]
            if btype == "heading":
                level = min(block["level"], 6)
                parts.append(
                    f"<h{level}>{self._escape_html(block['text'])}</h{level}>"
                )
            elif btype == "paragraph":
                parts.append(f"<p>{self._inline_to_html(block['text'])}</p>")
            elif btype == "list_item":
                tag = "ol" if block["ordered"] else "ul"
                parts.append(
                    f"<{tag}><li>{self._inline_to_html(block['text'])}</li></{tag}>"
                )
            elif btype == "table":
                parts.append(
                    '<table style="border-collapse:collapse;width:100%;'
                    'margin:8px 0;">'
                )
                parts.append("<thead><tr>")
                for h in block["headers"]:
                    parts.append(
                        '<th style="border:1px solid #bfbfbf;padding:4px 8px;'
                        'background:#f0f0f0;font-weight:700;text-align:left;">'
                        f"{self._escape_html(h)}</th>"
                    )
                parts.append("</tr></thead><tbody>")
                for row in block["rows"]:
                    parts.append("<tr>")
                    for c in row:
                        parts.append(
                            '<td style="border:1px solid #bfbfbf;padding:4px 8px;">'
                            f"{self._escape_html(c)}</td>"
                        )
                    parts.append("</tr>")
                parts.append("</tbody></table>")
            elif btype == "quote":
                parts.append(
                    '<blockquote style="border-left:2px solid #bfbfbf;'
                    'margin:8px 0;padding:4px 12px;color:#595959;">'
                    f"{self._inline_to_html(block['text'])}</blockquote>"
                )
        return parts

    def _inline_to_html(self, text: str) -> str:
        """把 **加粗** 转成 <b>（先 escape HTML）."""
        from app.services.report_markdown import _inline_to_html

        return _inline_to_html(text)

    @staticmethod
    def _escape_html(text: str) -> str:
        """转义 HTML 特殊字符（< > & " '）."""
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    @staticmethod
    def _cover_meta_lines(task: Task) -> list[str]:
        """封面元信息块（等宽对齐风格文本，缺项跳过，不报错）.

        返回若干行纯文本：
          被审查人：{name} · {id} · {department}（缺项跳过）
          审查周期：{start} ~ {end}
          生成日期：{now YYYY-MM-DD}
          任务编号：#{task.id}
        """
        from datetime import date

        lines: list[str] = []
        # 被审查人（拼 name/id/department，缺项跳过）。
        subject_parts = [
            p for p in (
                task.employee_name,
                task.employee_id,
                task.department,
            ) if p
        ]
        if subject_parts:
            lines.append("被审查人：" + " · ".join(subject_parts))
        # 审查周期。
        start_str = task.audit_start.strftime("%Y-%m-%d") if task.audit_start else None
        end_str = task.audit_end.strftime("%Y-%m-%d") if task.audit_end else None
        if start_str or end_str:
            lines.append(
                f"审查周期：{start_str or '—'} ~ {end_str or '—'}"
            )
        # 生成日期。
        lines.append("生成日期：" + date.today().strftime("%Y-%m-%d"))
        # 任务编号。
        lines.append(f"任务编号：#{task.id}")
        return lines

    @staticmethod
    def _docx_horizontal_rule(doc) -> None:
        """在 docx 文档当前位置插入一条横线（黑色、单色、底边框）."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 0.75pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    @staticmethod
    def _docx_insert_toc_field(doc) -> None:
        """在 docx 文档当前位置插入原生 TOC 域（``TOC \\o "1-2"``）.

        Word 打开时若域未计算，用户需「右键→更新域」才显示页码（python-docx
        无法预渲染域值，PRD 接受此限制）。结构：fldChar begin → instrText →
        fldChar separate → placeholder → fldChar end。
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        p = doc.add_paragraph()
        run = p.add_run()
        # begin.
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar_begin)
        # instrText.
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-2" \\h \\z \\u'
        run._r.append(instr)
        # separate.
        fldChar_sep = OxmlElement("w:fldChar")
        fldChar_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fldChar_sep)
        # placeholder（域值未计算时的占位文本）。
        placeholder = OxmlElement("w:t")
        placeholder.text = "右键此处选择「更新域」以生成目录。"
        run._r.append(placeholder)
        # end.
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar_end)

    async def _load_task(self, db: AsyncSession, task_id: int) -> Task:
        result = await db.execute(select(Task).where(Task.id == task_id))
        task = result.scalar_one_or_none()
        if not task:
            raise ValueError("Task not found")
        return task

    async def _load_review(self, db: AsyncSession, task_id: int, review_id: int | None) -> Review | None:
        query = select(Review).where(Review.task_id == task_id)
        if review_id is not None:
            query = query.where(Review.id == review_id)
        else:
            query = query.order_by(Review.created_at.desc())
        result = await db.execute(query.limit(1))
        return result.scalar_one_or_none()

    async def _load_matches(self, db: AsyncSession, review_id: int | None) -> list[ReviewMatch]:
        if review_id is None:
            return []
        result = await db.execute(
            select(ReviewMatch)
            .where(ReviewMatch.review_id == review_id)
            .order_by(ReviewMatch.record_id.asc())
        )
        return list(result.scalars().all())

    def _write_records_sheet(
        self,
        ws,
        records: list[FlowRecord],
        matches: list[ReviewMatch],
    ) -> None:
        ws.title = "标准化流水"
        best_by_record = {match.record_id: match for match in matches}
        headers = [
            "来源文件",
            "原始行号",
            "交易时间",
            "交易对手名",
            "交易对手账号",
            "金额",
            "摘要",
            "收支类型",
            "匹配用户",
            "匹配度",
            "匹配类型",
        ]
        self._write_header(ws, headers)
        for row_index, record in enumerate(records, 2):
            match = best_by_record.get(record.record_id)
            values: list[Any] = [
                record.source_file,
                record.original_row,
                record.transaction_time,
                record.counterparty_name,
                record.counterparty_account,
                record.amount,
                record.summary,
                record.transaction_type,
                match.customer_name if match else "",
                match.score if match else "",
                match.match_type if match else "",
            ]
            for col_index, value in enumerate(values, 1):
                ws.cell(row=row_index, column=col_index, value=value)

    def _write_matches_sheet(self, wb, matches: list[ReviewMatch]) -> None:
        ws = wb.create_sheet("匹配详情")
        headers = [
            "流水记录ID",
            "匹配用户",
            "匹配度",
            "匹配类型",
            "来源文件",
            "交易时间",
            "交易对手名",
            "交易对手账号",
            "金额",
            "摘要",
        ]
        self._write_header(ws, headers)
        for row_index, match in enumerate(matches, 2):
            values = [
                match.record_id,
                match.customer_name,
                match.score,
                match.match_type,
                match.source_file or "",
                match.transaction_time or "",
                match.counterparty_name or "",
                match.counterparty_account or "",
                match.amount or "",
                match.summary or "",
            ]
            for col_index, value in enumerate(values, 1):
                ws.cell(row=row_index, column=col_index, value=value)

    @staticmethod
    def _write_header(ws, headers: list[str]) -> None:
        fill = PatternFill(start_color="1F2937", end_color="1F2937", fill_type="solid")
        font = Font(color="FFFFFF", bold=True)
        for col_index, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_index, value=header)
            cell.fill = fill
            cell.font = font
            ws.column_dimensions[cell.column_letter].width = max(12, min(len(header) + 8, 30))
        ws.freeze_panes = "A2"

    @staticmethod
    def _match_to_dict(match: ReviewMatch) -> dict[str, Any]:
        return {
            "record_id": match.record_id,
            "customer_name": match.customer_name,
            "match_type": match.match_type,
            "score": match.score,
            "counterparty_name": match.counterparty_name,
            "counterparty_account": match.counterparty_account,
            "source_file": match.source_file,
            "transaction_time": match.transaction_time,
            "amount": match.amount,
            "summary": match.summary,
        }

    @staticmethod
    def _skill_markdown(task: Task, match_count: int) -> str:
        return (
            "# 员工客户流水审查 Skills\n\n"
            f"- 任务：{task.title}\n"
            f"- 命中记录：{match_count}\n\n"
            "使用 `current_task/review_result.json` 回答审查命中、证据明细和复核建议问题。\n"
        )


async def load_export(db: AsyncSession, export_id: int) -> ExportFile | None:
    """Load export artifact by id."""
    result = await db.execute(select(ExportFile).where(ExportFile.id == export_id))
    return result.scalar_one_or_none()

# -*- coding: utf-8 -*-
"""报告 markdown 约束子集渲染器（md → docx + md → pdf）.

06-28-report-fusion-word-cover Phase 1：把章节 markdown（report_chapter_builder
拼装的确定性模板 / 后续 LLM agent 产出）正确渲染成 Word/PDF 富文本，不再把
``##``/``|``/``-`` 当字面纯文本显示。

约束子集（prd §一.2，硬约束，渲染器只认这些）：
  * ``## `` / ``### `` 节标题（``# `` 一级留给章标题，渲染器不主动产出一级）
  * 段落：空行分隔的连续非空行合并
  * 列表：``- `` / ``* `` 无序 + ``1. `` 有序
  * 表格：标准 markdown 表格（``| a | b |`` + ``|---|---|`` 分隔行）
  * 引用：``> ``
  * 内联：``**加粗**``
  * 容错：认不全的当段落（不崩，不丢字）

单色硬底线（web-pages-design 单色设计系统）：表格表头灰底（``#f0f0f0``）加粗，
无红黄绿；引用灰字缩进。全站只用 9 级明度。

不引重依赖（纯手写行级状态机），只依赖已有 python-docx / reportlab。
"""

import re
from typing import Any

# 单色设计系统明度标尺（web-pages-design.md）。
GRAY_HEADER_BG = "#f0f0f0"  # 表头灰底（ink-300 对应的画布灰，单色表头）。
GRAY_QUOTE_TEXT = "#595959"  # 引用灰字（中灰辅文）。
BLACK_TEXT = "#1f1f1f"  # 正文深灰。


# ---------------------------------------------------------------------------
# 解析：parse_markdown_blocks(md) -> list[dict]
# ---------------------------------------------------------------------------


def _is_table_separator(line: str) -> bool:
    """识别 markdown 表格分隔行（``|---|---|`` / ``|:--:|:-|--:|`）.

    规则：去掉首尾 ``|`` 后，剩余单元每段都只含 ``-`` / ``:`` / 空格，且至少含
    一个 ``-``。畸形行（如 ``| abc |`` 非分隔行）返 False。
    """
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    # 去掉首尾 | 后按 | 切。
    inner = stripped.strip("|")
    cells = [c.strip() for c in inner.split("|")]
    if not cells:
        return False
    for c in cells:
        # 每段只允许 - / : / 空格，且至少一个 -.
        if not c:
            continue
        if not re.fullmatch(r"[\-:\s]+", c):
            return False
        if "-" not in c:
            return False
    return True


def _parse_table_row(line: str) -> list[str]:
    """解析一行 markdown 表格行 → 单元格列表（去首尾空格、处理转义 ``\\|``）.

    用占位符保护转义的 ``\\|``（``\\|`` → ``\\x00`` → split 后还原），避免
    被当成单元格分隔符。
    """
    stripped = line.strip()
    # 去掉首尾的 |（允许行首/行尾无 |）。
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    # 用占位符保护转义的 \|（NUL 字符不会出现在正常 markdown 里）。
    protected = stripped.replace("\\|", "\x00")
    cells = protected.split("|")
    return [c.strip().replace("\x00", "|") for c in cells]


def _split_heading(line: str) -> tuple[int, str] | None:
    """识别 ``##``/``###`` 等标题行 → (level, text)，否则 None.

    ``#`` 一级也识别（level=1，但渲染器约定不主动产出一级——章标题已由
    export 的 ``add_heading(level=1)`` 处理）。
    """
    m = re.match(r"^(#{1,6})\s+(.*)$", line.strip())
    if not m:
        return None
    level = len(m.group(1))
    text = m.group(2).strip()
    if not text:
        return None
    return level, text


def _split_list_item(line: str) -> tuple[bool, str] | None:
    """识别列表行 → (ordered, text)，否则 None.

    无序：``- `` / ``* `` / ``+ ``；有序：``1. `` / ``12. ``.
    """
    m = re.match(r"^\s*([-*+])\s+(.*)$", line)
    if m:
        return False, m.group(2).strip()
    m = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
    if m:
        return True, m.group(2).strip()
    return None


def _split_quote(line: str) -> str | None:
    """识别引用行 ``> xxx`` → text，否则 None."""
    m = re.match(r"^>\s?(.*)$", line.strip())
    if m:
        return m.group(1).strip()
    return None


def parse_markdown_blocks(md: str) -> list[dict[str, Any]]:
    """行级状态机解析约束子集 markdown → 块列表.

    每块是 dict，``type`` 取值：
      * ``{"type": "heading", "level": int, "text": str}``
      * ``{"type": "paragraph", "text": str}``
      * ``{"type": "list_item", "ordered": bool, "text": str}``
      * ``{"type": "table", "headers": [...], "rows": [[...], ...]}``
      * ``{"type": "quote", "text": str}``

    解析规则（容错）：
      * 空行分隔块（连续空行 = 块边界）。
      * 表格：遇到 ``|`` 行，如果下一行是分隔行则识别为表格；否则当作段落。
        连续的 ``|`` 行归入同一表格直到空行/非表格行。
      * 列表：连续的列表行归入同一组（渲染时各自独立块，但连续识别）。
      * 认不全的行 → 段落（连续非空行合并）。
    """
    if not md or not md.strip():
        return []

    lines = md.splitlines()
    blocks: list[dict[str, Any]] = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 空行跳过（块边界）。
        if not stripped:
            i += 1
            continue

        # 标题。
        heading = _split_heading(line)
        if heading:
            level, text = heading
            blocks.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # 引用：连续引用行合并为一段。
        quote_first = _split_quote(line)
        if quote_first is not None:
            quote_parts: list[str] = [quote_first]
            i += 1
            while i < n:
                q = _split_quote(lines[i])
                if q is None:
                    break
                if q:  # 非空引用行。
                    quote_parts.append(q)
                i += 1
            blocks.append({"type": "quote", "text": " ".join(quote_parts)})
            continue

        # 列表：连续列表行（连续识别，渲染各自独立块）。
        list_first = _split_list_item(line)
        if list_first is not None:
            while i < n:
                li = _split_list_item(lines[i])
                if li is None:
                    break
                ordered, text = li
                blocks.append({"type": "list_item", "ordered": ordered, "text": text})
                i += 1
            continue

        # 表格：当前行是 | 行 + 下一行是分隔行。
        if stripped.startswith("|") and i + 1 < n and _is_table_separator(lines[i + 1]):
            headers = _parse_table_row(line)
            i += 2  # 跳过表头行 + 分隔行。
            rows: list[list[str]] = []
            while i < n:
                row_line = lines[i]
                if not row_line.strip().startswith("|"):
                    break
                rows.append(_parse_table_row(row_line))
                i += 1
            blocks.append({"type": "table", "headers": headers, "rows": rows})
            continue

        # 段落：连续非空、非特殊块行合并。
        para_parts: list[str] = [stripped]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_stripped = nxt.strip()
            if not nxt_stripped:
                break
            # 遇到块边界（标题/引用/列表/表格）则段落结束。
            if (
                _split_heading(nxt)
                or _split_quote(nxt) is not None
                or _split_list_item(nxt)
                or (
                    nxt_stripped.startswith("|")
                    and i + 1 < n
                    and _is_table_separator(lines[i + 1])
                )
            ):
                break
            para_parts.append(nxt_stripped)
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_parts)})

    return blocks


# ---------------------------------------------------------------------------
# inline：**加粗** 解析
# ---------------------------------------------------------------------------

# ``**xxx**`` 加粗段（非贪婪，避免跨段误吞）。
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _parse_inline_spans(text: str) -> list[tuple[str, bool]]:
    """解析内联 ``**加粗**`` → [(text, bold), ...] spans 列表.

    用于 docx（拆 bold run + 普通 run）。
    """
    spans: list[tuple[str, bool]] = []
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            spans.append((text[pos:m.start()], False))
        spans.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        spans.append((text[pos:], False))
    if not spans:
        spans.append((text, False))
    return spans


def _inline_to_html(text: str) -> str:
    """把 ``**加粗**`` 转成 ``<b>加粗</b>``（先 escape HTML）.

    用于 reportlab Paragraph（支持基本 HTML 内联标签）。
    """
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    # 在 escape 后的文本上替换 **xxx** → <b>xxx</b>（** 在 escape 后不变）。
    return _BOLD_RE.sub(r"<b>\1</b>", escaped)


# ---------------------------------------------------------------------------
# docx 渲染：render_docx(doc, blocks)
# ---------------------------------------------------------------------------


def _docx_set_cell_shading(cell, fill_hex: str) -> None:
    """给 docx 单元格设灰底（w:shd fill）."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    # 移除已存在的 shd 避免重复。
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _docx_add_runs_with_bold(paragraph, text: str) -> None:
    """向 paragraph 加 run，``**bold**`` 拆成 bold run + 普通 run."""
    from docx.shared import Pt, RGBColor

    for span_text, bold in _parse_inline_spans(text):
        run = paragraph.add_run(span_text)
        run.bold = bold
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)


def render_docx(doc: Any, blocks: list[dict[str, Any]]) -> None:
    """把 blocks 渲染进 python-docx Document.

    * heading → ``doc.add_heading(text, level)``（level 2/3；1 留给章标题）
    * paragraph → ``add_paragraph`` + ``**bold**`` 拆 run
    * list_item → ``add_paragraph(text, style="List Bullet"/"List Number")``
    * table → ``add_table`` + 表头灰底加粗
    * quote → 缩进段落 + 灰字
    """
    from docx.shared import Pt, RGBColor

    for block in blocks:
        btype = block["type"]

        if btype == "heading":
            # 一级留给章标题（export 的 add_heading(level=1)），渲染器不产出 1。
            level = block["level"]
            heading = doc.add_heading(block["text"], level=max(2, level))
            # 单色：标题黑色。
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        elif btype == "paragraph":
            p = doc.add_paragraph()
            _docx_add_runs_with_bold(p, block["text"])

        elif btype == "list_item":
            style = "List Number" if block["ordered"] else "List Bullet"
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                # 模板无该 style 时回退普通段落 + 手动前缀。
                p = doc.add_paragraph()
                prefix = "- " if not block["ordered"] else "1. "
                run = p.add_run(prefix)
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
            _docx_add_runs_with_bold(p, block["text"])

        elif btype == "table":
            headers = block["headers"]
            rows = block["rows"]
            col_count = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=col_count)
            # 表头：加粗 + 灰底。
            for col_idx, header_text in enumerate(headers):
                cell = table.cell(0, col_idx)
                p = cell.paragraphs[0]
                # 清空单元格已有内容（python-docx 默认空段落）.
                for existing_run in list(p.runs):
                    existing_run.text = ""
                run = p.add_run(header_text)
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
                _docx_set_cell_shading(cell, "f0f0f0")
            # 数据行。
            for row_idx, row in enumerate(rows, 1):
                for col_idx in range(col_count):
                    cell = table.cell(row_idx, col_idx)
                    p = cell.paragraphs[0]
                    for existing_run in list(p.runs):
                        existing_run.text = ""
                    cell_text = row[col_idx] if col_idx < len(row) else ""
                    _docx_add_runs_with_bold(p, cell_text)

        elif btype == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            for span_text, bold in _parse_inline_spans(block["text"]):
                run = p.add_run(span_text)
                run.bold = bold
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)  # 中灰引用字.


# ---------------------------------------------------------------------------
# pdf 渲染：render_pdf(flowables, blocks, styles)
# ---------------------------------------------------------------------------


def render_pdf(
    flowables: list[Any],
    blocks: list[dict[str, Any]],
    styles: dict[str, Any] | None = None,
) -> None:
    """把 blocks 渲染成 reportlab flowables，追加进 flowables 列表.

    Args:
        flowables: 待追加的 reportlab flowable 列表（调用方持有）。
        blocks: ``parse_markdown_blocks`` 产物。
        styles: 可选样式 dict，键：``"h2"``/``"h3"``/``"body"``/``"quote"``，
            值为 ``ParagraphStyle``；缺省用简单默认样式。

    Styles 缺省时用极简单色样式（黑色标题、深灰正文、中灰引用）。
    """
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import (
        ListItem,
        ListFlowable,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )

    if styles is None:
        base = getSampleStyleSheet()
        styles = {
            "h2": ParagraphStyle(
                "MdH2",
                parent=base["Heading2"],
                fontName="Helvetica-Bold",
                fontSize=14,
                leading=18,
                spaceBefore=10,
                spaceAfter=4,
                textColor=HexColor(BLACK_TEXT),
            ),
            "h3": ParagraphStyle(
                "MdH3",
                parent=base["Heading3"],
                fontName="Helvetica-Bold",
                fontSize=12,
                leading=16,
                spaceBefore=8,
                spaceAfter=3,
                textColor=HexColor(BLACK_TEXT),
            ),
            "body": ParagraphStyle(
                "MdBody",
                parent=base["BodyText"],
                fontName="Helvetica",
                fontSize=10.5,
                leading=16,
                spaceAfter=6,
                textColor=HexColor(BLACK_TEXT),
                alignment=TA_LEFT,
            ),
            "quote": ParagraphStyle(
                "MdQuote",
                parent=base["BodyText"],
                fontName="Helvetica",
                fontSize=10.5,
                leading=15,
                leftIndent=18,
                spaceAfter=6,
                textColor=HexColor(GRAY_QUOTE_TEXT),
            ),
        }

    # 列表项缓冲（连续同序/同无序合并成一个 ListFlowable）。
    pending_list: list[tuple[bool, str]] = []  # [(ordered, text)].

    def flush_list() -> None:
        """把缓冲的列表项刷成一个 ListFlowable."""
        if not pending_list:
            return
        ordered = pending_list[0][0]
        items = []
        for _, text in pending_list:
            items.append(
                ListItem(
                    Paragraph(_inline_to_html(text), styles["body"]),
                    leftIndent=18,
                )
            )
        flowables.append(
            ListFlowable(
                items,
                bulletType="1" if ordered else "bullet",
                start="1" if ordered else None,
                leftIndent=18,
            )
        )
        flowables.append(Spacer(1, 3))
        pending_list.clear()

    for block in blocks:
        btype = block["type"]

        # 遇到非列表块时先刷掉缓冲列表。
        if btype != "list_item":
            flush_list()

        if btype == "heading":
            level = block["level"]
            style = styles["h2"] if level <= 2 else styles["h3"]
            flowables.append(Paragraph(_inline_to_html(block["text"]), style))

        elif btype == "paragraph":
            flowables.append(
                Paragraph(_inline_to_html(block["text"]), styles["body"])
            )

        elif btype == "list_item":
            pending_list.append((block["ordered"], block["text"]))

        elif btype == "table":
            headers = block["headers"]
            rows = block["rows"]
            # reportlab 表格：首行表头加粗 + 灰底。
            data = [[Paragraph(_inline_to_html(h), styles["body"]) for h in headers]]
            for row in rows:
                data.append(
                    [
                        Paragraph(
                            _inline_to_html(row[c] if c < len(row) else ""),
                            styles["body"],
                        )
                        for c in range(len(headers))
                    ]
                )
            col_widths = None
            n_cols = len(headers)
            if n_cols:
                # 等宽列（页面宽由 reportlab 自动 fit）。
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.units import mm

                page_w = A4[0] - 40 * mm  # 左右各 20mm 边距.
                col_widths = [page_w / n_cols] * n_cols
            table = Table(data, colWidths=col_widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        # 表头：加粗（HTML <b> 已加）+ 灰底 + 黑字.
                        ("BACKGROUND", (0, 0), (-1, 0), HexColor(GRAY_HEADER_BG)),
                        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor(BLACK_TEXT)),
                        # 全表网格线（单色浅灰）.
                        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#bfbfbf")),
                        # 内边距.
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        # 垂直对齐顶部.
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            flowables.append(table)
            flowables.append(Spacer(1, 4))

        elif btype == "quote":
            flowables.append(
                Paragraph(_inline_to_html(block["text"]), styles["quote"])
            )

    # 收尾：刷掉最后缓冲的列表。
    flush_list()

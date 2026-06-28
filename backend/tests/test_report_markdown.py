# -*- coding: utf-8 -*-
"""报告 markdown 约束子集渲染器 tests（06-28-report-fusion-word-cover Phase 1）.

Covers (per prd §3 验证):
* parse_markdown_blocks：各子集（##/###、段落、-/* 列表、1. 有序、表格带分隔行、
  > 引用、**加粗**内联）+ 边界（空串、畸形表格缺分隔行→当段落、混合块）。
* render_docx smoke：能生成 docx 不崩；关键块（heading/table/list）产物结构断言。
* render_pdf smoke：能生成 pdf 不崩；关键块 flowable 类型断言。
"""

from pathlib import Path

import pytest

from app.services.report_markdown import (
    parse_markdown_blocks,
    render_docx,
    render_pdf,
)


# ---------------------------------------------------------------------------
# parse_markdown_blocks：各子集
# ---------------------------------------------------------------------------


def test_parse_empty_string():
    """空串 → 空列表."""
    assert parse_markdown_blocks("") == []
    assert parse_markdown_blocks("   \n  \n") == []


def test_parse_heading_levels():
    """## / ### / # 标题正确识别 level."""
    blocks = parse_markdown_blocks("## 二级标题\n\n### 三级标题\n\n# 一级标题")
    assert len(blocks) == 3
    assert blocks[0] == {"type": "heading", "level": 2, "text": "二级标题"}
    assert blocks[1] == {"type": "heading", "level": 3, "text": "三级标题"}
    assert blocks[2] == {"type": "heading", "level": 1, "text": "一级标题"}


def test_parse_paragraph_multiline_merge():
    """连续非空行合并成一个段落（空行分隔）."""
    md = "第一行 第二行\n同一段落第二行\n\n第二段落"
    blocks = parse_markdown_blocks(md)
    assert len(blocks) == 2
    assert blocks[0]["type"] == "paragraph"
    assert "第一行" in blocks[0]["text"]
    assert "同一段落第二行" in blocks[0]["text"]
    assert blocks[1]["type"] == "paragraph"
    assert blocks[1]["text"] == "第二段落"


def test_parse_unordered_list_dash_and_star():
    """- 和 * 都识别为无序列表."""
    md = "- 第一项\n- 第二项\n* 第三项"
    blocks = parse_markdown_blocks(md)
    assert all(b["type"] == "list_item" for b in blocks)
    assert all(b["ordered"] is False for b in blocks)
    assert blocks[0]["text"] == "第一项"
    assert blocks[1]["text"] == "第二项"
    assert blocks[2]["text"] == "第三项"


def test_parse_ordered_list():
    """1. 2. 识别为有序列表."""
    md = "1. 第一步\n2. 第二步\n3. 第三步"
    blocks = parse_markdown_blocks(md)
    assert all(b["type"] == "list_item" for b in blocks)
    assert all(b["ordered"] is True for b in blocks)
    assert blocks[0]["text"] == "第一步"
    assert blocks[2]["text"] == "第三步"


def test_parse_table_with_separator():
    """标准 markdown 表格（带分隔行）识别为 table 块."""
    md = (
        "| 关键词 | 金额 | 风险 |\n"
        "| --- | --- | --- |\n"
        "| 词A | 5000 | 高 |\n"
        "| 词B | 3000 | 中 |"
    )
    blocks = parse_markdown_blocks(md)
    assert len(blocks) == 1
    table = blocks[0]
    assert table["type"] == "table"
    assert table["headers"] == ["关键词", "金额", "风险"]
    assert table["rows"] == [["词A", "5000", "高"], ["词B", "3000", "中"]]


def test_parse_table_pipe_escape():
    """表格单元内的转义 \\| 还原成 |."""
    md = "| 片段 |\n| --- |\n| 含\\|竖线 |"
    blocks = parse_markdown_blocks(md)
    assert blocks[0]["type"] == "table"
    assert blocks[0]["rows"][0][0] == "含|竖线"


def test_parse_quote_block():
    """> 引用行识别为 quote 块."""
    md = "> 这是引用\n> 第二行引用"
    blocks = parse_markdown_blocks(md)
    assert len(blocks) == 1
    assert blocks[0]["type"] == "quote"
    assert "这是引用" in blocks[0]["text"]
    assert "第二行引用" in blocks[0]["text"]


def test_parse_bold_inline_in_paragraph():
    """**加粗** 标记保留在段落 text 里（渲染层处理）."""
    md = "普通文字 **加粗内容** 普通"
    blocks = parse_markdown_blocks(md)
    assert blocks[0]["type"] == "paragraph"
    assert "**加粗内容**" in blocks[0]["text"]


# ---------------------------------------------------------------------------
# parse_markdown_blocks：边界与容错
# ---------------------------------------------------------------------------


def test_parse_malformed_table_no_separator_becomes_paragraph():
    """畸形表格（缺分隔行）→ 当段落（不崩）."""
    md = "| a | b |\n| c | d |"
    blocks = parse_markdown_blocks(md)
    # 没有 |---| 分隔行 → 不识别为 table，当段落.
    assert all(b["type"] == "paragraph" for b in blocks)
    assert len(blocks) >= 1


def test_parse_mixed_blocks():
    """混合块（标题 + 段落 + 列表 + 表格 + 引用）顺序正确."""
    md = (
        "## 概述\n"
        "\n"
        "这是概述段落。\n"
        "\n"
        "- 要点一\n"
        "- 要点二\n"
        "\n"
        "| 表头A | 表头B |\n"
        "| --- | --- |\n"
        "| x | y |\n"
        "\n"
        "> 这是引用\n"
    )
    blocks = parse_markdown_blocks(md)
    types = [b["type"] for b in blocks]
    assert types[0] == "heading"
    assert "paragraph" in types
    assert "list_item" in types
    assert "table" in types
    assert "quote" in types
    # heading 在最前.
    assert blocks[0] == {"type": "heading", "level": 2, "text": "概述"}


def test_parse_unknown_lines_become_paragraph():
    """不认的行（如纯文本）→ 段落（容错）."""
    md = "随便一行文字\n另一行"
    blocks = parse_markdown_blocks(md)
    assert all(b["type"] == "paragraph" for b in blocks)


def test_parse_keyword_review_chapter_template():
    """真实模板章（_build_keyword_review 产物）能完整解析."""
    md = (
        "## 关键词审查\n"
        "\n"
        "已确认命中总数：2\n"
        "\n"
        "按风险等级分组（灰阶递进，高 > 中 > 低）：\n"
        "- 高风险：1 项\n"
        "- 中风险：1 项\n"
        "\n"
        "已确认命中明细：\n"
        "\n"
        "| 关键词 | 对手方 | 金额 | 命中字段 | 命中片段 | 风险等级 | 匹配类型 |\n"
        "| --- | --- | --- | --- | --- | --- | --- |\n"
        "| 卡/词 | 某公司 | 5000 | 对手方 | 片段 | 高 | 精确匹配 |\n"
        "\n"
        "（仅已确认(confirmed)的命中计入。）"
    )
    blocks = parse_markdown_blocks(md)
    types = [b["type"] for b in blocks]
    assert "heading" in types
    assert "paragraph" in types
    assert "list_item" in types
    assert "table" in types
    # 表格结构完整.
    table = [b for b in blocks if b["type"] == "table"][0]
    assert len(table["headers"]) == 7
    assert len(table["rows"]) == 1


# ---------------------------------------------------------------------------
# render_docx smoke：能生成 docx + 关键块结构断言
# ---------------------------------------------------------------------------


def test_render_docx_heading_and_paragraph(tmp_path: Path):
    """docx 渲染含 heading + paragraph."""
    from docx import Document

    doc = Document()
    blocks = [
        {"type": "heading", "level": 2, "text": "标题"},
        {"type": "paragraph", "text": "正文段落"},
    ]
    render_docx(doc, blocks)
    path = tmp_path / "test.docx"
    doc.save(str(path))
    assert path.exists() and path.stat().st_size > 0

    # 重新读验结构.
    doc2 = Document(str(path))
    paragraphs = doc2.paragraphs
    # 第一个 paragraph 应是 heading 样式.
    found_heading = any("标题" in p.text for p in paragraphs)
    found_body = any("正文段落" in p.text for p in paragraphs)
    assert found_heading
    assert found_body


def test_render_docx_table_with_header_shading(tmp_path: Path):
    """docx 渲染表格 + 表头灰底（w:shd fill=f0f0f0）."""
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document()
    blocks = [
        {
            "type": "table",
            "headers": ["列A", "列B"],
            "rows": [["1", "2"], ["3", "4"]],
        }
    ]
    render_docx(doc, blocks)
    path = tmp_path / "test_table.docx"
    doc.save(str(path))

    doc2 = Document(str(path))
    assert len(doc2.tables) == 1
    table = doc2.tables[0]
    # 3 行（1 表头 + 2 数据）× 2 列.
    assert len(table.rows) == 3
    assert len(table.columns) == 2
    # 表头单元有灰底 shd.
    header_cell = table.cell(0, 0)
    tc_pr = header_cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill")) == "f0f0f0"
    # 表头文字加粗.
    assert table.cell(0, 0).paragraphs[0].runs[0].bold is True


def test_render_docx_list_bullet_and_number(tmp_path: Path):
    """docx 渲染列表（无序 + 有序）."""
    from docx import Document

    doc = Document()
    blocks = [
        {"type": "list_item", "ordered": False, "text": "无序项"},
        {"type": "list_item", "ordered": True, "text": "有序项"},
    ]
    render_docx(doc, blocks)
    path = tmp_path / "test_list.docx"
    doc.save(str(path))

    doc2 = Document(str(path))
    texts = [p.text for p in doc2.paragraphs]
    assert any("无序项" in t for t in texts)
    assert any("有序项" in t for t in texts)


def test_render_docx_bold_inline(tmp_path: Path):
    """docx 渲染 **加粗** → bold run."""
    from docx import Document

    doc = Document()
    blocks = [{"type": "paragraph", "text": "前 **加粗** 后"}]
    render_docx(doc, blocks)
    path = tmp_path / "test_bold.docx"
    doc.save(str(path))

    doc2 = Document(str(path))
    # 找含「加粗」的段落.
    target = None
    for p in doc2.paragraphs:
        if "加粗" in p.text:
            target = p
            break
    assert target is not None
    # 应有 bold run.
    bold_runs = [r for r in target.runs if r.bold]
    assert len(bold_runs) >= 1
    assert any(r.text == "加粗" for r in bold_runs)


# ---------------------------------------------------------------------------
# render_pdf smoke：能生成 pdf + 关键块 flowable 断言
# ---------------------------------------------------------------------------


def test_render_pdf_heading_and_body(tmp_path: Path):
    """pdf 渲染含 heading + paragraph."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph

    flowables = []
    blocks = [
        {"type": "heading", "level": 2, "text": "节标题"},
        {"type": "paragraph", "text": "正文内容"},
    ]
    render_pdf(flowables, blocks)
    # flowables 含 Paragraph.
    paragraphs = [f for f in flowables if isinstance(f, Paragraph)]
    assert len(paragraphs) >= 2

    # 真正生成 pdf 文件不崩.
    path = tmp_path / "test.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    doc.build(flowables)
    assert path.exists() and path.stat().st_size > 0


def test_render_pdf_table_with_header_background(tmp_path: Path):
    """pdf 渲染表格 + 表头灰底（TableStyle BACKGROUND）."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Table

    flowables = []
    blocks = [
        {
            "type": "table",
            "headers": ["列A", "列B"],
            "rows": [["1", "2"]],
        }
    ]
    render_pdf(flowables, blocks)
    tables = [f for f in flowables if isinstance(f, Table)]
    assert len(tables) == 1
    # 真正生成 pdf 文件不崩（表头灰底由 TableStyle BACKGROUND 设置，smoke 验证）.
    path = tmp_path / "test_table.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    doc.build(flowables)
    assert path.exists() and path.stat().st_size > 0


def test_render_pdf_list_flowable(tmp_path: Path):
    """pdf 渲染列表 → ListFlowable."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, ListFlowable

    flowables = []
    blocks = [
        {"type": "list_item", "ordered": False, "text": "无序项"},
        {"type": "list_item", "ordered": False, "text": "第二项"},
    ]
    render_pdf(flowables, blocks)
    list_flowables = [f for f in flowables if isinstance(f, ListFlowable)]
    # 连续同序列表项合并成一个 ListFlowable.
    assert len(list_flowables) == 1

    path = tmp_path / "test_list.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    doc.build(flowables)
    assert path.exists() and path.stat().st_size > 0


def test_render_pdf_quote(tmp_path: Path):
    """pdf 渲染引用 → Paragraph 缩进灰字."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph

    flowables = []
    blocks = [{"type": "quote", "text": "引用文字"}]
    render_pdf(flowables, blocks)
    paragraphs = [f for f in flowables if isinstance(f, Paragraph)]
    assert len(paragraphs) >= 1

    path = tmp_path / "test_quote.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=A4)
    doc.build(flowables)
    assert path.exists() and path.stat().st_size > 0


# ---------------------------------------------------------------------------
# 端到端：真实模板章 markdown → docx/pdf 不崩
# ---------------------------------------------------------------------------


def test_render_real_chapter_markdown_to_docx_and_pdf(tmp_path: Path):
    """真实模板章 markdown（含表格/列表/标题）完整渲染 docx + pdf 不崩."""
    md = (
        "## 数据范围\n"
        "\n"
        "- 文档数：3\n"
        "- 标准化记录数：150\n"
        "- 渠道分布：\n"
        "\n"
        "| 渠道 | 笔数 |\n"
        "| --- | --- |\n"
        "| 银行 | 100 |\n"
        "| 支付 | 50 |\n"
        "\n"
        "数据范围以清洗标准化提交锁定的 standard 快照为准。"
    )
    blocks = parse_markdown_blocks(md)
    # 解析结构正确.
    assert any(b["type"] == "table" for b in blocks)
    assert any(b["type"] == "list_item" for b in blocks)

    # docx 不崩.
    from docx import Document

    doc = Document()
    render_docx(doc, blocks)
    docx_path = tmp_path / "chapter.docx"
    doc.save(str(docx_path))
    assert docx_path.stat().st_size > 0

    # pdf 不崩.
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate

    flowables = []
    render_pdf(flowables, blocks)
    pdf_path = tmp_path / "chapter.pdf"
    pdf_doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
    pdf_doc.build(flowables)
    assert pdf_path.stat().st_size > 0

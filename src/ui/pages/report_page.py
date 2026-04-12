# -*- coding: utf-8 -*-
"""
Report page - LLM-generated summary and detailed report for current review task.
"""

from pathlib import Path

from PyQt5.QtCore import QObject, QThread, pyqtSignal
from PyQt5.QtWidgets import QWidget, QLabel, QPushButton, QHBoxLayout, QVBoxLayout, QFileDialog, QMessageBox, QTextEdit

from ..styles import COLORS
from ..widgets import Card, StatCardRow
from ...config import get_config
from ...export_flows import SkillsExporter
from ...llm import AuditAgent


class ReportWorker(QObject):
    """后台生成审查报告。"""

    finished = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(self, task_title: str, task_id: str, result):
        super().__init__()
        self.task_title = task_title
        self.task_id = task_id
        self.result = result

    def run(self) -> None:
        try:
            agent = AuditAgent()
            report = agent.generate_report(self.task_title, self.task_id, self.result)
            self.finished.emit(report)
        except Exception as exc:
            self.error.emit(str(exc))


class ReportPage(QWidget):
    """当前审查任务的报告页面。"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.config = get_config()
        self.result = None
        self.task_title = ""
        self.task_id = ""
        self._worker = None
        self._worker_thread = None
        self.generated_report = ""
        self._setup_ui()

    def _setup_ui(self) -> None:
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(20)

        title = QLabel("审查报告")
        title.setObjectName("title")
        layout.addWidget(title)

        self.subtitle = QLabel("当前任务暂无审查结果")
        self.subtitle.setStyleSheet(f"color: {COLORS['text_secondary']}; font-size: 13px;")
        layout.addWidget(self.subtitle)

        self.stats_row = StatCardRow()
        self.stats_row.add_stat("审查文档", "0", "份", COLORS["primary"])
        self.stats_row.add_stat("流水笔数", "0", "笔", COLORS["success"])
        self.stats_row.add_stat("匹配条数", "0", "条", COLORS["warning"])
        layout.addWidget(self.stats_row)

        summary_card = Card()
        summary_card.add_title("大模型报告")

        top_actions = QHBoxLayout()
        self.status_label = QLabel("请先完成一次审查。")
        self.status_label.setStyleSheet(f"color: {COLORS['text_secondary']}; font-size: 12px;")
        top_actions.addWidget(self.status_label)
        top_actions.addStretch()

        self.regenerate_btn = QPushButton("重新生成")
        self.regenerate_btn.setObjectName("secondary_btn")
        self.regenerate_btn.clicked.connect(self._generate_report)
        top_actions.addWidget(self.regenerate_btn)
        summary_card.layout.addLayout(top_actions)

        self.summary_output = QTextEdit()
        self.summary_output.setReadOnly(True)
        self.summary_output.setMinimumHeight(420)
        summary_card.layout.addWidget(self.summary_output)
        layout.addWidget(summary_card)

        action_layout = QHBoxLayout()
        action_layout.addStretch()

        self.export_btn = QPushButton("导出报告文档")
        self.export_btn.setObjectName("secondary_btn")
        self.export_btn.clicked.connect(self._export_report)
        action_layout.addWidget(self.export_btn)

        self.export_skill_btn = QPushButton("导出 Skills 包")
        self.export_skill_btn.setObjectName("secondary_btn")
        self.export_skill_btn.clicked.connect(self._export_skills_bundle)
        action_layout.addWidget(self.export_skill_btn)

        layout.addLayout(action_layout)

    def set_report_result(self, result, task_title: str = "", task_id: str = "") -> None:
        self.result = result
        self.task_title = str(task_title or "")
        self.task_id = str(task_id or "")

        if not result:
            self.clear()
            return

        card0 = self.stats_row.get_card(0)
        card1 = self.stats_row.get_card(1)
        card2 = self.stats_row.get_card(2)
        if card0:
            card0.set_value(str(self._count_reviewed_documents(result)))
        if card1:
            card1.set_value(str(self._count_reviewed_flows(result)))
        if card2:
            card2.set_value(str(result.total_matches))

        title_part = self.task_title or self.task_id or "当前任务"
        self.subtitle.setText(f"{title_part} 的审查报告")
        self.summary_output.setPlainText("正在调用大模型生成报告...")
        self._generate_report()

    def _generate_report(self) -> None:
        if not self.result:
            QMessageBox.warning(self, "无数据", "请先完成一次审查")
            return

        self._cleanup_worker()
        self.status_label.setText("正在调用大模型生成报告...")
        self.summary_output.setPlainText("正在调用大模型生成报告，请稍候...")
        self.regenerate_btn.setEnabled(False)

        self._worker_thread = QThread()
        self._worker = ReportWorker(self.task_title, self.task_id, self.result)
        self._worker.moveToThread(self._worker_thread)
        self._worker_thread.started.connect(self._worker.run)
        self._worker.finished.connect(self._on_report_finished)
        self._worker.error.connect(self._on_report_error)
        self._worker_thread.start()

    def _on_report_finished(self, report_text: str) -> None:
        self.generated_report = report_text
        self.summary_output.setPlainText(report_text)
        self.status_label.setText("大模型报告已生成")
        self.regenerate_btn.setEnabled(True)
        self._cleanup_worker()

    def _on_report_error(self, message: str) -> None:
        self.generated_report = ""
        self.summary_output.setPlainText(f"报告生成失败：\n{message}")
        self.status_label.setText("大模型报告生成失败")
        self.regenerate_btn.setEnabled(True)
        self._cleanup_worker()

    def _cleanup_worker(self) -> None:
        if self._worker_thread:
            self._worker_thread.quit()
            self._worker_thread.wait(3000)
            self._worker_thread = None
        self._worker = None

    def _export_report(self) -> None:
        if not self.result:
            QMessageBox.warning(self, "无数据", "当前没有可导出的报告")
            return

        default_name = f"审查报告_{self.task_id or self.result.review_id}.docx"
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出审查报告",
            str(self.config.reports_folder / default_name),
            "Word Files (*.docx)",
        )
        if not file_path:
            return
        if not file_path.lower().endswith(".docx"):
            file_path = f"{file_path}.docx"

        try:
            report_text = self.generated_report or self.summary_output.toPlainText()
            from docx import Document

            document = Document()
            document.add_heading(self.task_title or "审查报告", level=0)

            summary_rows = [
                f"任务编号：{self.task_id or self.result.review_id}",
                f"审查时间：{self.result.review_time}",
                f"审查文档数：{self._count_reviewed_documents(self.result)} 份",
                f"流水笔数：{self._count_reviewed_flows(self.result)} 笔",
                f"客户数量：{self.result.total_customers} 个",
                f"命中客户：{self.result.matched_customers} 个",
                f"匹配记录：{self.result.total_matches} 条",
                f"审查金额：{self.result.total_amount_formatted}",
            ]
            if self.result.writeback_error:
                summary_rows.append(f"写回状态：失败（{self.result.writeback_error}）")
            else:
                summary_rows.append("写回状态：成功")

            for line in summary_rows:
                document.add_paragraph(line)

            document.add_heading("报告正文", level=1)
            lines = (report_text or "").splitlines() or [""]
            for line in lines:
                text = str(line).rstrip()
                if not text:
                    document.add_paragraph("")
                    continue
                if text.startswith(("一、", "二、", "三、", "四、")):
                    document.add_heading(text, level=2)
                else:
                    document.add_paragraph(text)

            if self.result.matches:
                document.add_heading("附录：匹配明细", level=1)
                table = document.add_table(rows=1, cols=7)
                table.style = "Table Grid"
                headers = ["匹配用户", "来源文件", "交易时间", "对手名", "对手账号", "金额", "摘要"]
                for col_idx, header in enumerate(headers):
                    table.rows[0].cells[col_idx].text = header
                for match in self.result.matches:
                    cells = table.add_row().cells
                    cells[0].text = str(match.customer_name or "")
                    cells[1].text = str(Path(match.source_file).name if match.source_file else "")
                    cells[2].text = str(match.transaction_time or "")
                    cells[3].text = str(match.counterparty_name or "")
                    cells[4].text = str(match.counterparty_account or "")
                    cells[5].text = str(match.amount or "")
                    cells[6].text = str(match.summary or "")

            Path(file_path).parent.mkdir(parents=True, exist_ok=True)
            document.save(file_path)
            QMessageBox.information(self, "导出成功", f"审查报告已导出到:\n{file_path}")
        except Exception as exc:
            QMessageBox.warning(self, "导出失败", f"导出报告时出错:\n{exc}")

    def _export_skills_bundle(self) -> None:
        if not self.result:
            QMessageBox.warning(self, "无数据", "当前没有可导出的 Skills 包")
            return

        default_name = f"审查Skills包_{self.task_id or self.result.review_id}.zip"
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "导出 Skills 包",
            str(self.config.reports_folder / default_name),
            "Zip Files (*.zip)",
        )
        if not file_path:
            return
        if not file_path.lower().endswith(".zip"):
            file_path = f"{file_path}.zip"

        try:
            exporter = SkillsExporter(self.config)
            output_path = exporter.export_bundle(
                self.result,
                task_title=self.task_title,
                task_id=self.task_id,
                report_text=self.generated_report or self.summary_output.toPlainText(),
                output_path=file_path,
            )
            QMessageBox.information(
                self,
                "导出成功",
                "Skills 包已导出到:\n"
                f"{output_path}\n\n"
                "其中包含当前任务资料、标准化流水、固定工作流说明，以及历史审查目录。",
            )
        except Exception as exc:
            QMessageBox.warning(self, "导出失败", f"导出 Skills 包时出错:\n{exc}")

    def clear(self) -> None:
        self._cleanup_worker()
        self.result = None
        self.task_title = ""
        self.task_id = ""
        self.generated_report = ""
        self.subtitle.setText("当前任务暂无审查结果")
        self.status_label.setText("请先完成一次审查。")
        self.summary_output.setPlainText("")
        card0 = self.stats_row.get_card(0)
        card1 = self.stats_row.get_card(1)
        card2 = self.stats_row.get_card(2)
        if card0:
            card0.set_value("0")
        if card1:
            card1.set_value("0")
        if card2:
            card2.set_value("0")
        self.regenerate_btn.setEnabled(True)

    @staticmethod
    def _count_reviewed_documents(result) -> int:
        try:
            from openpyxl import load_workbook

            wb = load_workbook(result.flow_excel_path, read_only=True, data_only=True)
            try:
                ws = wb.active
                rows_iter = ws.iter_rows(values_only=True)
                headers = next(rows_iter, None)
                if not headers:
                    return 0

                source_idx = None
                for idx, header in enumerate(headers):
                    if str(header or "").strip() == "来源文件":
                        source_idx = idx
                        break

                if source_idx is None:
                    return 0

                files = set()
                for row in rows_iter:
                    value = row[source_idx] if source_idx < len(row) else ""
                    text = str(value or "").strip()
                    if text:
                        files.add(text)
                return len(files)
            finally:
                wb.close()
        except Exception:
            return 0

    @staticmethod
    def _count_reviewed_flows(result) -> int:
        try:
            from openpyxl import load_workbook

            wb = load_workbook(result.flow_excel_path, read_only=True, data_only=True)
            try:
                ws = wb.active
                return max(ws.max_row - 1, 0)
            finally:
                wb.close()
        except Exception:
            return 0
